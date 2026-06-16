from docx import Document
from typing import Dict, List, Any
import re


class DocumentParser:
    def __init__(self):
        pass
    
    def parse_docx(self, file_path: str) -> Dict[str, Any]:
        doc = Document(file_path)
        
        result = {
            "title": self._extract_title(doc),
            "sections": self._extract_sections(doc),
            "tables": self._extract_tables(doc),
            "parameters": self._extract_parameters(doc),
            "clauses": self._extract_clauses(doc)
        }
        
        return result
    
    def _extract_title(self, doc: Document) -> str:
        if doc.paragraphs:
            for para in doc.paragraphs[:5]:
                if para.text.strip():
                    return para.text.strip()
        return ""
    
    def _extract_sections(self, doc: Document) -> List[Dict[str, Any]]:
        sections = []
        current_section = None
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            if para.style.name.startswith('Heading'):
                level = int(para.style.name.replace('Heading', ''))
                if current_section:
                    sections.append(current_section)
                current_section = {
                    "level": level,
                    "title": text,
                    "content": [],
                    "subsections": []
                }
            elif current_section:
                current_section["content"].append(text)
        
        if current_section:
            sections.append(current_section)
        
        return sections
    
    def _extract_tables(self, doc: Document) -> List[Dict[str, Any]]:
        tables = []
        
        for i, table in enumerate(doc.tables):
            table_data = {
                "index": i,
                "rows": []
            }
            
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data["rows"].append(row_data)
            
            tables.append(table_data)
        
        return tables
    
    def _extract_parameters(self, doc: Document) -> List[Dict[str, Any]]:
        parameters = []
        param_pattern = re.compile(r'[\u4e00-\u9fa5a-zA-Z]+[：:]\s*[\d.]+')
        
        for para in doc.paragraphs:
            text = para.text.strip()
            matches = param_pattern.findall(text)
            
            for match in matches:
                parts = re.split(r'[：:]', match)
                if len(parts) == 2:
                    parameters.append({
                        "name": parts[0].strip(),
                        "value": parts[1].strip(),
                        "source": text
                    })
        
        for table in doc.tables:
            for row in table.rows:
                if len(row.cells) >= 2:
                    name = row.cells[0].text.strip()
                    value = row.cells[1].text.strip()
                    if name and value and any(c.isdigit() for c in value):
                        parameters.append({
                            "name": name,
                            "value": value,
                            "source": "table"
                        })
        
        return parameters
    
    def _extract_clauses(self, doc: Document) -> List[Dict[str, Any]]:
        clauses = []
        clause_pattern = re.compile(r'^[\d.]+\s+')
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if clause_pattern.match(text):
                clauses.append({
                    "number": clause_pattern.match(text).group().strip(),
                    "content": text
                })
        
        return clauses
    
    def get_full_text(self, file_path: str) -> str:
        doc = Document(file_path)
        full_text = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
        
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                full_text.append(row_text)
        
        return "\n".join(full_text)
